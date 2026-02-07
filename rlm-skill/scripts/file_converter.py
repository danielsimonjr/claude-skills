#!/usr/bin/env python3
"""
file_converter.py - Detect and convert various file types to text for RLM processing.

Supports:
- PDF files (.pdf)
- Word documents (.docx)
- Text files (.txt, .md, .rst)
- Code files (.py, .js, .ts, .java, .c, .cpp, .go, .rs, etc.)
- Data files (.json, .xml, .csv, .yaml, .yml)
- HTML files (.html, .htm)
- Archives containing text files (.zip, .tar, .tar.gz)

Usage:
    python file_converter.py input_file [output_file]
    python file_converter.py document.pdf                    # outputs to stdout
    python file_converter.py document.pdf extracted.txt      # outputs to file

Requires (installed automatically if missing):
- pdfplumber (for PDFs)
- python-docx (for Word documents)
- beautifulsoup4 (for HTML)
"""

import sys
import os
import json
import subprocess
from pathlib import Path
from typing import Optional, Tuple
import tempfile
import shutil


# File extension mappings
TEXT_EXTENSIONS = {
    '.txt', '.md', '.rst', '.text', '.log', '.ini', '.cfg', '.conf',
    '.py', '.js', '.ts', '.jsx', '.tsx', '.java', '.c', '.cpp', '.h', '.hpp',
    '.go', '.rs', '.rb', '.php', '.swift', '.kt', '.scala', '.r', '.m',
    '.sh', '.bash', '.zsh', '.fish', '.ps1', '.bat', '.cmd',
    '.sql', '.graphql', '.proto',
    '.json', '.xml', '.yaml', '.yml', '.toml', '.csv', '.tsv',
    '.html', '.htm', '.css', '.scss', '.sass', '.less',
    '.makefile', '.dockerfile', '.gitignore', '.env',
}

CODE_EXTENSIONS = {
    '.py', '.js', '.ts', '.jsx', '.tsx', '.java', '.c', '.cpp', '.h', '.hpp',
    '.go', '.rs', '.rb', '.php', '.swift', '.kt', '.scala', '.cs', '.vb',
    '.sh', '.bash', '.ps1', '.sql',
}


def detect_file_type(filepath: str) -> str:
    """
    Detect the type of file based on extension and content.
    
    Returns one of: 'pdf', 'docx', 'text', 'code', 'json', 'xml', 'html', 
                    'csv', 'archive', 'binary', 'unknown'
    """
    path = Path(filepath)
    ext = path.suffix.lower()
    
    # Check by extension first
    if ext == '.pdf':
        return 'pdf'
    elif ext == '.docx':
        return 'docx'
    elif ext == '.doc':
        return 'doc_legacy'
    elif ext in {'.json', '.jsonl'}:
        return 'json'
    elif ext == '.xml':
        return 'xml'
    elif ext in {'.html', '.htm'}:
        return 'html'
    elif ext in {'.csv', '.tsv'}:
        return 'csv'
    elif ext in {'.yaml', '.yml'}:
        return 'yaml'
    elif ext in {'.zip', '.tar', '.gz', '.tgz', '.tar.gz'}:
        return 'archive'
    elif ext in CODE_EXTENSIONS:
        return 'code'
    elif ext in TEXT_EXTENSIONS:
        return 'text'
    
    # Try to detect by content
    try:
        with open(filepath, 'rb') as f:
            header = f.read(16)
        
        # PDF magic bytes
        if header.startswith(b'%PDF'):
            return 'pdf'
        # ZIP (also docx)
        if header.startswith(b'PK\x03\x04'):
            # Check if it's a docx
            if ext == '.docx' or b'word/' in open(filepath, 'rb').read(1000):
                return 'docx'
            return 'archive'
        # Check if it's readable text
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                f.read(1000)
            return 'text'
        except UnicodeDecodeError:
            return 'binary'
            
    except Exception:
        return 'unknown'


def install_package(package: str) -> bool:
    """Install a Python package if not already installed."""
    try:
        __import__(package.replace('-', '_').split('[')[0])
        return True
    except ImportError:
        print(f"Installing {package}...", file=sys.stderr)
        result = subprocess.run(
            [sys.executable, '-m', 'pip', 'install', package, '--break-system-packages', '-q'],
            capture_output=True
        )
        return result.returncode == 0


MAX_PDF_SIZE = 500_000_000  # 500MB


def extract_pdf(filepath: str) -> str:
    """Extract text from PDF file."""
    file_size = os.path.getsize(filepath)
    if file_size > MAX_PDF_SIZE:
        raise ValueError(
            f"PDF too large ({file_size / 1_000_000:.0f}MB). "
            f"Max allowed: {MAX_PDF_SIZE / 1_000_000:.0f}MB."
        )
    # Try pdfplumber first (better quality)
    if install_package('pdfplumber'):
        import pdfplumber
        text_parts = []
        with pdfplumber.open(filepath) as pdf:
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"--- Page {i+1} ---\n{page_text}")
        if text_parts:
            return '\n\n'.join(text_parts)
    
    # Fallback to pdftotext command line tool
    try:
        result = subprocess.run(
            ['pdftotext', '-layout', filepath, '-'],
            capture_output=True, text=True,
            encoding='utf-8', errors='replace'
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
    except FileNotFoundError:
        pass
    
    # Last resort: PyPDF2
    if install_package('PyPDF2'):
        import PyPDF2
        text_parts = []
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    text_parts.append(f"--- Page {i+1} ---\n{text}")
        return '\n\n'.join(text_parts)
    
    raise RuntimeError("Could not extract PDF text. Install pdfplumber: pip install pdfplumber")


def extract_docx(filepath: str) -> str:
    """Extract text from Word document."""
    if not install_package('python-docx'):
        raise RuntimeError("Install python-docx: pip install python-docx")
    
    from docx import Document
    doc = Document(filepath)
    
    text_parts = []
    
    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            # Check if it's a heading
            if para.style.name.startswith('Heading'):
                level = para.style.name.replace('Heading ', '')
                text_parts.append(f"{'#' * int(level) if level.isdigit() else '#'} {para.text}")
            else:
                text_parts.append(para.text)
    
    # Extract tables
    for table in doc.tables:
        table_text = []
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells)
            table_text.append(row_text)
        if table_text:
            text_parts.append('\n[TABLE]\n' + '\n'.join(table_text) + '\n[/TABLE]')
    
    return '\n\n'.join(text_parts)


def extract_html(filepath: str) -> str:
    """Extract text from HTML file."""
    if install_package('beautifulsoup4'):
        from bs4 import BeautifulSoup
        with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
            soup = BeautifulSoup(f.read(), 'html.parser')
        
        # Remove script and style elements
        for element in soup(['script', 'style', 'nav', 'footer', 'header']):
            element.decompose()
        
        return soup.get_text(separator='\n', strip=True)
    else:
        # Fallback: basic regex stripping
        import re
        with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
            content = f.read()
        # Remove tags
        text = re.sub(r'<[^>]+>', ' ', content)
        # Clean whitespace
        text = re.sub(r'\s+', ' ', text)
        return text.strip()


def _is_safe_tar_member(member: 'tarfile.TarInfo', dest_dir: str) -> bool:
    """Check if a tar member extracts safely within the destination directory."""
    member_path = os.path.realpath(os.path.join(dest_dir, member.name))
    return member_path.startswith(os.path.realpath(dest_dir) + os.sep) or member_path == os.path.realpath(dest_dir)


def extract_archive(filepath: str) -> str:
    """Extract and concatenate text files from archive."""
    import zipfile
    import tarfile

    text_parts = []
    temp_dir = tempfile.mkdtemp()

    try:
        # Extract archive
        if filepath.endswith('.zip'):
            with zipfile.ZipFile(filepath, 'r') as zf:
                zf.extractall(temp_dir)
        elif filepath.endswith(('.tar', '.tar.gz', '.tgz', '.gz')):
            with tarfile.open(filepath, 'r:*') as tf:
                safe_members = [m for m in tf.getmembers() if _is_safe_tar_member(m, temp_dir)]
                tf.extractall(temp_dir, members=safe_members)
        
        # Process extracted files
        for root, dirs, files in os.walk(temp_dir):
            # Skip hidden directories and common non-text dirs
            dirs[:] = [d for d in dirs if not d.startswith('.') and d not in {'node_modules', '__pycache__', 'venv', '.git'}]
            
            for file in sorted(files):
                if file.startswith('.'):
                    continue
                    
                file_path = os.path.join(root, file)
                rel_path = os.path.relpath(file_path, temp_dir)
                
                file_type = detect_file_type(file_path)
                if file_type in ('text', 'code', 'json', 'xml', 'html', 'csv', 'yaml'):
                    try:
                        content = convert_to_text(file_path)
                        text_parts.append(f"=== FILE: {rel_path} ===\n{content}")
                    except Exception as e:
                        text_parts.append(f"=== FILE: {rel_path} ===\n[Error reading file: {e}]")
        
        return '\n\n'.join(text_parts)
        
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)


def extract_json(filepath: str) -> str:
    """Convert JSON to readable text format."""
    with open(filepath, 'r', encoding='utf-8') as f:
        # Handle JSON Lines format
        if filepath.endswith('.jsonl'):
            lines = f.readlines()
            objects = [json.loads(line) for line in lines if line.strip()]
            return '\n---\n'.join(json.dumps(obj, indent=2) for obj in objects)
        else:
            data = json.load(f)
            return json.dumps(data, indent=2)


def extract_text(filepath: str) -> str:
    """Read plain text file."""
    with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
        return f.read()


def convert_to_text(filepath: str) -> str:
    """
    Convert any supported file type to text.
    
    Args:
        filepath: Path to the input file
        
    Returns:
        Extracted text content
        
    Raises:
        ValueError: If file type is not supported
        RuntimeError: If extraction fails
    """
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"File not found: {filepath}")
    
    file_type = detect_file_type(filepath)
    
    extractors = {
        'pdf': extract_pdf,
        'docx': extract_docx,
        'html': extract_html,
        'archive': extract_archive,
        'json': extract_json,
        'text': extract_text,
        'code': extract_text,
        'xml': extract_text,
        'csv': extract_text,
        'yaml': extract_text,
    }
    
    if file_type in extractors:
        return extractors[file_type](filepath)
    elif file_type == 'doc_legacy':
        raise ValueError(
            "Legacy .doc format not supported. Please convert to .docx or PDF first.\n"
            "You can use LibreOffice: libreoffice --convert-to docx document.doc"
        )
    elif file_type == 'binary':
        raise ValueError(f"Binary file detected. Cannot extract text from: {filepath}")
    else:
        raise ValueError(f"Unsupported file type: {file_type} for {filepath}")


def get_file_info(filepath: str) -> dict:
    """Get information about a file."""
    path = Path(filepath)
    file_type = detect_file_type(filepath)
    
    info = {
        'path': str(path.absolute()),
        'name': path.name,
        'extension': path.suffix,
        'detected_type': file_type,
        'size_bytes': path.stat().st_size,
        'size_human': format_size(path.stat().st_size),
    }
    
    return info


def format_size(size_bytes: int) -> str:
    """Format file size in human-readable format."""
    for unit in ['B', 'KB', 'MB', 'GB']:
        if size_bytes < 1024:
            return f"{size_bytes:.1f} {unit}"
        size_bytes /= 1024
    return f"{size_bytes:.1f} TB"


def main():
    import argparse
    
    parser = argparse.ArgumentParser(
        description='Convert various file types to text for RLM processing',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Supported formats:
  Documents:  .pdf, .docx, .txt, .md, .rst
  Code:       .py, .js, .ts, .java, .c, .cpp, .go, .rs, .rb, .php, etc.
  Data:       .json, .jsonl, .xml, .yaml, .yml, .csv, .tsv
  Web:        .html, .htm
  Archives:   .zip, .tar, .tar.gz, .tgz

Examples:
  python file_converter.py document.pdf
  python file_converter.py report.docx output.txt
  python file_converter.py codebase.zip extracted.txt --info
        """
    )
    
    parser.add_argument('input', help='Input file path')
    parser.add_argument('output', nargs='?', help='Output file path (default: stdout)')
    parser.add_argument('--info', '-i', action='store_true', help='Show file info only')
    
    args = parser.parse_args()
    
    if not os.path.exists(args.input):
        print(f"Error: File not found: {args.input}", file=sys.stderr)
        sys.exit(1)
    
    if args.info:
        info = get_file_info(args.input)
        print(f"File: {info['name']}")
        print(f"Type: {info['detected_type']}")
        print(f"Size: {info['size_human']}")
        print(f"Path: {info['path']}")
        return
    
    try:
        text = convert_to_text(args.input)
        
        if args.output:
            with open(args.output, 'w', encoding='utf-8') as f:
                f.write(text)
            print(f"Extracted {len(text):,} characters to {args.output}", file=sys.stderr)
        else:
            print(text)
            
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
